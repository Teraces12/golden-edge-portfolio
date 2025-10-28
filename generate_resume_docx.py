from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Create document
doc = Document()

# Set narrow margins for one-page format
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

# Helper function to add spacing
def set_space_after(paragraph, points):
    p_format = paragraph.paragraph_format
    p_format.space_after = Pt(points)
    p_format.space_before = Pt(0)
    p_format.line_spacing = 1.0

# Header with name and contact
header = doc.add_paragraph()
header_format = header.paragraph_format
header_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
header_format.space_after = Pt(2)
header_format.space_before = Pt(0)
run = header.add_run("LEBEDE NGARTERA")
run.bold = True
run.font.size = Pt(14)

# Title and contact on same line
subheader = doc.add_paragraph()
subheader_format = subheader.paragraph_format
subheader_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subheader_format.space_after = Pt(4)
subheader_format.space_before = Pt(0)
run = subheader.add_run("Senior ML Engineer | AI Research & Infrastructure | Upper Darby, PA | +1 (267) 961 1931 | ngarteralebede12@gmail.com | GitHub")
run.font.size = Pt(9)

# Professional Summary - Compact
summary_title = doc.add_paragraph()
summary_title_run = summary_title.add_run("PROFESSIONAL SUMMARY")
summary_title_run.bold = True
summary_title_run.font.size = Pt(10)
set_space_after(summary_title, 2)

summary = doc.add_paragraph()
summary_text = "Ph.D. in Applied Mathematics with 7+ years engineering scalable deep learning systems and distributed ML training. Expert in PyTorch (5+ yrs), Python (7+ yrs), and large-scale data pipelines. Proven ability to design, implement, and optimize highly-scalable ML frameworks for research and production. Track record: 93.1% accuracy, 44.8% improvement, 300x optimization. 9+ peer-reviewed publications. Passionate about translating cutting-edge research into reproducible, production-ready code."
summary.add_run(summary_text)
summary.paragraph_format.space_after = Pt(3)
summary.paragraph_format.space_before = Pt(0)
summary.paragraph_format.line_spacing = 1.0
for run in summary.runs:
    run.font.size = Pt(9)

# Core Skills - Compact
skills_title = doc.add_paragraph()
skills_title_run = skills_title.add_run("CORE SKILLS & EXPERIENCE")
skills_title_run.bold = True
skills_title_run.font.size = Pt(10)
set_space_after(skills_title, 2)

skills_items = [
    "Python (7+ yrs) • PyTorch (5+ yrs) • TensorFlow • C/C++ • R • SQL",
    "Distributed ML Training • Data Pipeline Architecture • PyTorch DataLoader Optimization",
    "Deep Learning Infrastructure • Algorithms & Optimization • Large-Scale Dataset Processing",
    "Azure • AWS • GCP • Docker • Git • Performance Metrics (F1, BLEU, ROUGE, AUC-ROC)"
]

for skill_item in skills_items:
    skill = doc.add_paragraph(skill_item, style='List Bullet')
    skill.paragraph_format.space_after = Pt(1)
    skill.paragraph_format.space_before = Pt(0)
    skill.paragraph_format.line_spacing = 1.0
    for run in skill.runs:
        run.font.size = Pt(9)

# Professional Experience
exp_title = doc.add_paragraph()
exp_title_run = exp_title.add_run("PROFESSIONAL EXPERIENCE")
exp_title_run.bold = True
exp_title_run.font.size = Pt(10)
set_space_after(exp_title, 2)

# Experience 1
exp1_title = doc.add_paragraph()
exp1_run = exp1_title.add_run("Senior ML Engineer & AI Researcher")
exp1_run.bold = True
exp1_run.font.size = Pt(9)
exp1_org = exp1_title.add_run(" | TeraSystemsAI, Independent Research | 2019 – Present")
exp1_org.font.size = Pt(9)
exp1_title.paragraph_format.space_after = Pt(1)
set_space_after(exp1_title, 1)

exp1_bullets = [
    "Engineered scalable PyTorch-based ML systems for production deployment; custom training loops, optimization techniques, 300x speedup",
    "Architected large-scale data pipelines and PyTorch DataLoader strategies; processed enterprise-scale documents (200-300 pages)",
    "Built and open-sourced reproducible research code; Bayesian RAG platform: 93.1% accuracy, 44.8% improvement over baseline",
    "Implemented ML algorithms, A/B testing frameworks, statistical validation (F1, BLEU, ROUGE); 9+ peer-reviewed publications"
]

for bullet in exp1_bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.size = Pt(9)

# Experience 2
exp2_title = doc.add_paragraph()
exp2_run = exp2_title.add_run("Data Scientist & Assistant Lecturer")
exp2_run.bold = True
exp2_run.font.size = Pt(9)
exp2_org = exp2_title.add_run(" | University of N'djamena | 2013 – 2018")
exp2_org.font.size = Pt(9)
set_space_after(exp2_title, 1)

exp2_bullets = [
    "Designed and implemented ML models for forecasting and quantitative analysis",
    "Taught AI, Machine Learning, Statistics to undergraduates and graduates; mentored research methodology"
]

for bullet in exp2_bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        run.font.size = Pt(9)

# Education & Certifications
edu_title = doc.add_paragraph()
edu_title_run = edu_title.add_run("EDUCATION & CERTIFICATIONS")
edu_title_run.bold = True
edu_title_run.font.size = Pt(10)
set_space_after(edu_title, 1)

edu = doc.add_paragraph()
edu_text = "Ph.D. in Applied Mathematics, Cheikh Anta Diop University of Dakar (Analysis, Statistics & Applications) • AWS SageMaker ML Engineer • IBM Data Science Specialization • Natural Language Processing (DeepLearning.AI)"
edu.add_run(edu_text)
edu.paragraph_format.space_after = Pt(3)
edu.paragraph_format.space_before = Pt(0)
edu.paragraph_format.line_spacing = 1.0
for run in edu.runs:
    run.font.size = Pt(9)

# Key Achievements
ach_title = doc.add_paragraph()
ach_title_run = ach_title.add_run("KEY ACHIEVEMENTS")
ach_title_run.bold = True
ach_title_run.font.size = Pt(10)
set_space_after(ach_title, 1)

ach_items = [
    "93.1% accuracy, 44.8% improvement over baseline; 300x production optimization; 99.9% uptime",
    "9+ peer-reviewed publications (IEEE Access 2025, Frontiers 2025, AI/ML conferences)",
    "Open-source reproducible research frameworks; scalable ML systems for research and production",
    "Distributed ML infrastructure; large-scale dataset optimization; research-to-product pipeline"
]

for ach_item in ach_items:
    ach = doc.add_paragraph(ach_item, style='List Bullet')
    ach.paragraph_format.space_after = Pt(1)
    ach.paragraph_format.space_before = Pt(0)
    ach.paragraph_format.line_spacing = 1.0
    for run in ach.runs:
        run.font.size = Pt(9)

# Save document
output_path = r"c:\Users\ngart\OneDrive\Desktop\job_ready\Lebede_Ngartera_Resume_Meta_FAIR.docx"
doc.save(output_path)
print(f"✅ Resume created successfully: {output_path}")
